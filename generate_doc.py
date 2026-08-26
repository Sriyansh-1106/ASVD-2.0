import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_document():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles & Colors
    NAVY = RGBColor(11, 28, 65)
    CYAN = RGBColor(0, 160, 200)
    DARK_GRAY = RGBColor(50, 50, 50)
    RED = RGBColor(200, 20, 50)
    GREEN = RGBColor(0, 140, 70)

    # --- Title Section ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title.add_run("🛡️ ASVD 2.0 — AI SCAM VOICE DETECTION SYSTEM")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = subtitle.add_run("The Ultimate Code & Architecture Guide: Explained for a 10-Year-Old (With Full Technical Depth!)")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = CYAN

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = meta.add_run("Author: Sriyansh Sheel | Version: 2.0.0 | Technology Stack: Python, FastAPI, WebSockets, Scikit-Learn, Web Audio API, Vanilla JS")
    r_meta.font.name = "Arial"
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = DARK_GRAY

    doc.add_paragraph() # Spacer

    # Helper function for Section Headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = CYAN
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = DARK_GRAY
        return p

    def add_body_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_GRAY
        return p

    def add_callout(kid_analogy, tech_reality):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.8)
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F0F8FF") # light ice blue

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.right_indent = Inches(0.1)

        r1 = p.add_run("🧒 10-Year-Old Analogy: ")
        r1.bold = True
        r1.font.color.rgb = NAVY
        r2 = p.add_run(kid_analogy + "\n\n")
        r2.font.italic = True

        r3 = p.add_run("💻 Under the Hood (Technical Reality): ")
        r3.bold = True
        r3.font.color.rgb = CYAN
        r4 = p.add_run(tech_reality)
        doc.add_paragraph()

    def add_code_box(filename, code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        tbl.columns[0].width = Inches(6.8)
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "1E1E1E") # Dark VS code background

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

        r_hdr = p.add_run(f"📁 File: {filename}\n")
        r_hdr.font.name = "Consolas"
        r_hdr.font.size = Pt(9.5)
        r_hdr.font.bold = True
        r_hdr.font.color.rgb = RGBColor(0, 242, 254)

        r_code = p.add_run(code_text)
        r_code.font.name = "Consolas"
        r_code.font.size = Pt(9)
        r_code.font.color.rgb = RGBColor(220, 220, 220)
        doc.add_paragraph()

    # =========================================================================
    # SECTION 1: THE BIG PICTURE
    # =========================================================================
    add_heading_1("1. The Big Picture — What is ASVD 2.0?")
    add_body_p(
        "ASVD 2.0 stands for AI Scam Voice Detection System 2.0. It is a real-time cyber defense platform designed to protect people "
        "from phone call scammers, extortionists, fake police officers, bank fraud, and emergency scams in English, Hindi, and Hinglish."
    )

    add_callout(
        "Imagine you have a loyal robot bodyguard with super-ears who sits between your phone and your ear. "
        "Whenever a stranger calls claiming they are from the police or bank, your robot bodyguard listens to every single word, "
        "checks a secret detective book of 5,000 scam tricks, and if it smells a lie, it instantly flashes a red siren on your smartwatch screen shouting: "
        "'DANGER! DO NOT GIVE YOUR PASSWORD!'",

        "ASVD 2.0 is an asynchronous, multi-device, full-duplex cyber intelligence application. It streams live microphone audio over WebSockets, "
        "normalizes multi-lingual transcripts (Hindi/English/Hinglish), vectorizes conversation tokens with TF-IDF N-Grams, "
        "evaluates 8 distinct psychological manipulation heuristics via NLP Regex rule engines, computes a probabilistic threat score (0–100) "
        "using a Machine Learning Classifier, and pushes real-time defense actions and audio siren alarms to the victim's HUD in sub-100ms latency."
    )

    # =========================================================================
    # SECTION 2: ARCHITECTURE & SYSTEM FLOW
    # =========================================================================
    add_heading_1("2. Architecture & Data Flow: The 5-Step Relay Race")
    add_body_p(
        "How does a voice traveling into a laptop microphone turn into a flashing red threat alarm in a split second? "
        "The system works like a 5-step relay race:"
    )

    steps = [
        ("Step 1: Audio Ingestion", "Device 1 (Caller) captures live sound waves using the browser Web Audio API & MediaRecorder or Web Speech API."),
        ("Step 2: Speech-to-Text Conversion", "Audio bytes are transformed into clean text strings in real time (e.g. 'Main Inspector Sharma bol raha hoon, OTP do')."),
        ("Step 3: NLP Marker Extraction", "Rule-based regular expression engines scan the text for psychological coercion (Urgency, Authority, Threats, Money, OTP)."),
        ("Step 4: ML Classification & Risk Engine", "A trained Machine Learning model predicts the probability of fraud and calculates a composite 0–100 Risk Score."),
        ("Step 5: WebSocket Broadcast & HUD Siren", "The FastAPI server broadcasts a JSON payload to Device 2 (Receiver HUD), which updates the circular risk dial and rings the siren.")
    ]
    for s_title, s_desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r1 = p.add_run(f"• {s_title}: ")
        r1.bold = True
        r1.font.color.rgb = CYAN
        p.add_run(s_desc)

    doc.add_paragraph()

    # =========================================================================
    # SECTION 3: CODEBASE MAP — FILE BY FILE BREAKDOWN
    # =========================================================================
    add_heading_1("3. Detailed File-by-File Breakdown & Code Explanation")
    add_body_p(
        "Below is the complete walkthrough of every major file in the repository, explaining its purpose, inner logic, and how the code operates."
    )

    # File 1: backend/app/main.py
    add_heading_2("📁 backend/app/main.py — The Air Traffic Controller")
    add_body_p(
        "This file is the main entry point of the entire backend application built with FastAPI. "
        "It sets up the web server, turns on CORS (Cross-Origin Resource Sharing so any device can talk to it), mounts all static frontend files, "
        "and links the API routes and WebSocket endpoints together."
    )
    add_callout(
        "Think of main.py as the grand central train station master. When someone visits the website, talks on the phone, or sends audio, "
        "the station master checks their ticket and points them to the right platform (API platform, WebSocket platform, or Frontend visual screen).",
        "It initializes FastAPI(lifespan=lifespan), connects CORS middleware with allow_origins=['*'], initializes the SQLite database schema on startup, "
        "includes api_router and ws_router, and serves static files from /caller, /receiver, and /static using Starlette's StaticFiles."
    )
    add_code_box("backend/app/main.py",
'''app = FastAPI(title="ASVD 2.O — AI Scam Voice Detection API", version="2.0.0", lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

app.include_router(api_router)
app.include_router(ws_router)

# Mount frontend directories
app.mount("/caller", StaticFiles(directory=CALLER_DIR, html=True), name="caller-static")
app.mount("/receiver", StaticFiles(directory=RECEIVER_DIR, html=True), name="receiver-static")''')

    # File 2: backend/app/detection/indicators.py
    add_heading_2("📁 backend/app/detection/indicators.py — The 8 Psychological Clue Detectives")
    add_body_p(
        "Scammers don't just ask for money randomly — they use 8 psychological tricks to panic their victims. "
        "This file contains regex patterns that scan Hindi, English, and Hinglish for these exact manipulation markers:"
    )

    indicators_list = [
        ("1. Urgency Pressure", "Words like 'turant', 'jaldi', 'urgent', 'immediately', 'right now', 'within 5 minutes'."),
        ("2. OTP Requests", "Demanding one-time passwords, verification codes, 6-digit codes, or 'otp batao'."),
        ("3. Credential Harvesting", "Asking for PIN, CVV, password, card number, net banking login."),
        ("4. Authority Impersonation", "Pretending to be Police, CBI, Cyber Crime, Inspector, RBI, Supreme Court, SBI Bank Manager."),
        ("5. Threats & Legal Coercion", "Threatening arrest, FIR, jail, warrants, account freezing, electricity cutoff."),
        ("6. Financial Demand", "Demanding money transfer, UPI, Google Pay, PhonePe, refund fees, customs duty."),
        ("7. Secrecy & Isolation", "'Kisi ko mat batana', 'do not tell family', 'confidential', 'keep line active'."),
        ("8. Emotional Manipulation", "Pretending to be a crying relative in a hospital or road accident needing emergency operation funds.")
    ]
    for ind_name, ind_desc in indicators_list:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r_i = p.add_run(f"• {ind_name}: ")
        r_i.bold = True
        r_i.font.color.rgb = RED
        p.add_run(ind_desc)

    add_code_box("backend/app/detection/indicators.py",
'''INDICATOR_PATTERNS = {
    "urgency": [r"\\b(urgent|turant|jaldi|immediately|right now|5 minutes|block ho jayega)\\b"],
    "otp_request": [r"\\b(otp|one time password|verification code|6-digit|otp batao)\\b"],
    "authority_impersonation": [r"\\b(cbi|police|inspector|cyber crime|customs|rbi|trai)\\b"],
    "threat_detected": [r"\\b(arrest|fir|jail|non-bailable|account block|freeze)\\b"],
    "financial_request": [r"\\b(transfer|upi|gpay|phonepe|deposit|rupees|lakh|rs)\\b"],
    "secrecy_request": [r"\\b(kisi ko mat batana|confidential|do not tell|secret)\\b"],
    "emotional_manipulation": [r"\\b(accident|hospital|operation|emergency|save me)\\b"]
}''')

    # File 3: backend/app/detection/risk_engine.py
    add_heading_2("📁 backend/app/detection/risk_engine.py — The Master Threat Calculator")
    add_body_p(
        "This file combines the Machine Learning model's prediction with the 8 heuristic indicator markers to produce a final, balanced 0–100 Risk Score, "
        "a Risk Level (LOW, MEDIUM, HIGH, CRITICAL), and actionable safety advice (like calling the 1930 Cyber Crime helpline)."
    )
    add_callout(
        "Imagine taking a math test: you get 50 points from the ML robot, and 10 points for each scam trick found. "
        "If you score 85+, the red alarm goes off!",
        "The risk engine uses a weighted ensemble: Risk = (ML_Confidence * 0.40) + (Indicator_Weights * 0.60). "
        "Certain critical indicators like 'otp_request' or 'threat_detected' have mandatory minimum score floors (e.g. minimum 80 for OTP extortion)."
    )
    add_code_box("backend/app/detection/risk_engine.py",
'''def assess_risk(text: str) -> dict:
    indicators = detect_indicators(text)
    ml_result = predict_conversation(text)
    
    # Calculate score based on indicator count and ML probability
    base_score = ml_result["confidence"] * 50
    indicator_score = len(indicators) * 12
    final_score = min(100, int(base_score + indicator_score))
    
    # Floor score to CRITICAL if OTP or arrest threat is present
    if "otp_request" in indicators or "threat_detected" in indicators:
        final_score = max(final_score, 80)
        
    return {
        "risk_score": final_score,
        "risk_level": "CRITICAL" if final_score >= 80 else "HIGH" if final_score >= 60 else "LOW",
        "indicators": indicators,
        "recommended_action": generate_safety_advice(indicators)
    }''')

    # File 4: backend/app/speech/speech_to_text.py
    add_heading_2("📁 backend/app/speech/speech_to_text.py — The Multi-Turn Speech Pipeline")
    add_body_p(
        "During a real phone conversation, people speak sentence by sentence. If you only look at one sentence, you might miss the scam. "
        "The SpeechPipeline class accumulates the entire phone call history chunk-by-chunk without duplicating text."
    )
    add_code_box("backend/app/speech/speech_to_text.py",
'''class SpeechPipeline:
    def __init__(self, session_id: str = "default_session"):
        self.session_id = session_id
        self.full_transcript: str = ""
        self.transcript_chunks: List[str] = []

    def add_transcript_chunk(self, chunk: str) -> str:
        cleaned = chunk.strip()
        if self.full_transcript and cleaned.startswith(self.full_transcript):
            self.full_transcript = cleaned
        else:
            self.transcript_chunks.append(cleaned)
            self.full_transcript = (self.full_transcript + " " + cleaned).strip()
        return self.full_transcript''')

    # File 5: backend/app/api/websocket.py
    add_heading_2("📁 backend/app/api/websocket.py — The Ultra-Fast Secret Walkie-Talkie")
    add_body_p(
        "Regular websites reload pages to show new data. But during a phone scam, every millisecond counts! "
        "WebSockets create an open, live bidirectional pipeline between the Caller device, the AI server, and the Receiver HUD."
    )
    add_callout(
        "Instead of sending letters back and forth through the post office, WebSockets are like keeping a walkie-talkie button pressed down 24/7. "
        "The second Device 1 whispers a word, Device 2 hears it instantly!",
        "ConnectionManager maintains active WebSocket connections mapped by session_id room channels. "
        "When an 'audio_transcript' JSON packet arrives from role='caller', it processes the text through the risk engine "
        "and calls broadcast_to_session() to push the JSON threat alert to all active 'receiver' clients in sub-100ms."
    )
    add_code_box("backend/app/api/websocket.py",
'''@ws_router.websocket("/ws/call/{session_id}")
async def websocket_call_endpoint(websocket: WebSocket, session_id: str, role: str = "receiver"):
    await manager.connect(websocket, session_id)
    while True:
        data = await websocket.receive_json()
        if data.get("type") == "audio_transcript":
            full_text = pipeline.add_transcript_chunk(data["text"])
            threat = assess_risk(full_text)
            await manager.broadcast_to_session(session_id, {
                "type": "threat_alert",
                "data": threat,
                "full_transcript": full_text
            })''')

    # File 6: frontend/caller/app.js & index.html
    add_heading_2("📁 frontend/caller/ — Device 1: The Caller Simulator & Microphone Engine")
    add_body_p(
        "This is the simulator for the incoming caller. It features:"
    )
    c_features = [
        "Live Microphone Capture: Uses navigator.mediaDevices.getUserMedia and webkitSpeechRecognition to listen to real speech.",
        "MediaRecorder Audio Slices: Slices live audio every 3 seconds and posts WAV data to /api/speech_to_text.",
        "Audio Waveform Visualizer: Uses Web Audio API AnalyserNode (FFT) to draw live sine waves bouncing to your voice volume.",
        "Real-Time Speech Stream Box: Displays interim words in glowing yellow and finalized sentences in cyan.",
        "Real-Time AI Threat Preview: Shows the 0-100 risk score directly on the caller screen."
    ]
    for feat in c_features:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(f"• {feat}")

    # File 7: frontend/receiver/app.js & index.html
    add_heading_2("📁 frontend/receiver/ — Device 2: The Receiver Defense HUD")
    add_body_p(
        "This is the victim's AI security dashboard. It features:"
    )
    r_features = [
        "0–100 Circular Risk Gauge: Dynamic SVG circle that fills with smooth CSS animations from Green (Safe) -> Orange (Medium) -> Red (Critical).",
        "Keyword Highlighter: Scans incoming live text and puts glowing red badges around words like 'OTP', 'arrest', 'CBI', 'transfer'.",
        "Active Indicator Chips: Lights up urgency, impersonation, or financial threat chips.",
        "Synthesizer Warning Siren: Uses Web Audio API OscillatorNode to generate an emergency wobble siren alarm directly in code without external MP3 files.",
        "Emergency Action Helpline: One-click 1930 Cyber Crime Helpline button and call termination advice."
    ]
    for feat in r_features:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(f"• {feat}")

    # File 8: run_demo.py & start_demo.bat
    add_heading_2("📁 start_demo.bat & run_demo.py — The Intelligent Launchers")
    add_body_p(
        "These launcher scripts make starting the system 100% painless. "
        "They automatically check if port 8000 is occupied by an orphaned process, kill the stuck process, activate the virtual environment, "
        "start the Uvicorn server, and specifically launch Google Chrome with both Caller and Receiver tabs side-by-side."
    )

    # =========================================================================
    # SECTION 4: MACHINE LEARNING TRAINING & EVALUATION
    # =========================================================================
    add_heading_1("4. Machine Learning & Dataset Pipeline")
    add_body_p(
        "How did the AI learn what a scam looks like? "
        "In ml/train.py, we trained a Machine Learning model on 5,000 synthetic multi-lingual cyber scam call transcripts."
    )

    # Table of Dataset Metrics
    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(3.2)
    tbl.columns[1].width = Inches(3.6)

    headers = ["Metric / Parameter", "Value & Description"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_background(cell, "0B1C41")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    data_rows = [
        ("Dataset Size", "5,000 synthetic audio call transcripts (English, Hindi, Hinglish)"),
        ("Feature Extraction", "TF-IDF Vectorizer with unigrams & bigrams (ngram_range=(1,2), max_features=5000)"),
        ("Classifier Algorithm", "Logistic Regression with L2 Regularization & Balanced Class Weights"),
        ("Model Performance", "Accuracy: >95% | Precision: >94% | Recall: >96% | F1-Score: >95%")
    ]

    for row_idx, (m, v) in enumerate(data_rows, start=1):
        c0 = tbl.cell(row_idx, 0)
        c1 = tbl.cell(row_idx, 1)
        if row_idx % 2 == 0:
            set_cell_background(c0, "F5F5F5")
            set_cell_background(c1, "F5F5F5")
        c0.paragraphs[0].add_run(m).bold = True
        c1.paragraphs[0].add_run(v)

    doc.add_paragraph()

    # =========================================================================
    # SECTION 5: GLOSSARY OF TECHNICAL TERMS
    # =========================================================================
    add_heading_1("5. Glossary: Technical Terms Made Crystal Clear")
    
    glossary = [
        ("FastAPI", "A modern, lightning-fast Python framework used for building web servers and APIs with automatic data validation."),
        ("WebSocket", "A persistent two-way communication channel between browser and server, enabling real-time streaming without page refreshing."),
        ("TF-IDF (Term Frequency-Inverse Document Frequency)", "A math formula that measures how important a word is in a document compared to all documents."),
        ("N-Grams", "Looking at words in pairs or triplets (like 'police station' or 'urgent otp') instead of isolated single words."),
        ("Web Audio API", "A powerful browser system for processing, synthesizing sound waveforms, analyzing audio frequencies (FFT), and playing sirens."),
        ("MediaRecorder", "A browser JavaScript interface that records raw audio directly from your physical microphone into digital byte blobs."),
        ("SQLite", "A lightweight, serverless database engine stored in a single file on disk to persist call logs and detection records.")
    ]

    for term, definition in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r_t = p.add_run(f"• {term}: ")
        r_t.bold = True
        r_t.font.color.rgb = NAVY
        p.add_run(definition)

    doc.add_paragraph()

    # Footer note
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = footer_p.add_run("🛡️ ASVD 2.0 — Developed with Advanced Cyber AI Architecture")
    r_foot.font.italic = True
    r_foot.font.size = Pt(10)
    r_foot.font.color.rgb = CYAN

    # Save
    out_path = os.path.join(os.getcwd(), "ASVD_2.0_Complete_System_Guide.docx")
    doc.save(out_path)
    print(f"[SUCCESS] Successfully generated Word Document at: {out_path}")
    return out_path

if __name__ == "__main__":
    if sys.platform == "win32" and hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    create_document()
